from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract DOCX paragraphs and tables.")
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    document = Document(args.docx)
    print(f"PARAGRAPHS {len(document.paragraphs)}")
    print(f"TABLES {len(document.tables)}")

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"P{index} [{paragraph.style.name}]: {text}")

    for table_index, table in enumerate(document.tables):
        print(f"TABLE {table_index}")
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            print(f"R{row_index}: " + " | ".join(cells))


if __name__ == "__main__":
    main()
